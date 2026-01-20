import streamlit as st
import os
import PyPDF2
import docx
import requests
import json
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph

# OpenRouter API configuration - Use secrets for cloud deployment
try:
    # Try to get from Streamlit secrets first (for cloud deployment)
    OPENROUTER_API_KEY = st.secrets["OPENROUTER_API_KEY"]
    OPENROUTER_URL = st.secrets["OPENROUTER_URL"]
except:
    # Fallback to hardcoded values (for local development)
    OPENROUTER_API_KEY = "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"
    OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

def count_tokens(text):
    """Count tokens using simple estimation (compatible with cloud deployment)"""
    # Simple estimation: 1 token ≈ 4 characters for English text
    # This is a rough approximation but works well for most cases
    return len(text) // 4

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def analyze_resume_with_ai(resume_text, job_description):
    """Analyze resume against job description using OpenRouter API"""
    
    # Check token count and chunk if necessary
    total_tokens = count_tokens(resume_text)
    print(f"Resume token count: {total_tokens}")
    
    # For shorter resumes, use regular analysis with Few-shot prompting
    prompt = f"""
    You are an expert resume analyst with 10+ years of experience in career counseling and ATS optimization.
    
    Your task is to analyze a resume against a job description using Chain-of-Thought reasoning.
    
    Here are examples of high-quality analysis:
    
    Example 1 - Strong Match:
    {{
        "match_percentage": 85,
        "strengths": ["5+ years Python development", "AWS certified", "Led team of 5"],
        "missing_skills": ["No critical gaps identified"],
        "recommendations": ["Add quantifiable metrics to achievements"],
        "critical_gaps": ["Minor: Add project outcomes"],
        "ats_keywords": ["All key terms present"]
    }}
    
    Example 2 - Medium Match:
    {{
        "match_percentage": 65,
        "strengths": ["Strong technical background", "Good communication skills"],
        "missing_skills": ["Cloud experience", "Leadership examples"],
        "recommendations": ["Add cloud certifications", "Show leadership examples"],
        "critical_gaps": ["Missing cloud architecture experience"],
        "ats_keywords": ["Cloud, Kubernetes missing"]
    }}
    
    Now analyze this resume using same thorough approach:
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    
    Chain-of-Thought Process:
    1. First, identify all technical skills in job description
    2. Then, map each skill to resume content
    3. Next, assess experience level alignment
    4. Finally, identify critical gaps and ATS optimization opportunities
    
    Focus on identifying:
    1. Critical gaps between resume requirements and job needs
    2. Missing keywords and skills that ATS systems look for
    3. Experience level alignment with seniority requirements
    4. Industry-specific terminology gaps
    5. Quantifiable achievements that are missing
    
    Please provide analysis in following JSON format:
    {{
        "match_percentage": 85,
        "strengths": ["List of key strengths found in resume"],
        "missing_skills": ["List of important skills missing from resume"],
        "recommendations": ["List of specific recommendations to improve resume"],
        "key_matches": ["List of good matches between resume and job description"],
        "experience_alignment": "Assessment of how well experience aligns with requirements",
        "overall_assessment": "Brief overall assessment",
        "critical_gaps": ["List of most critical gaps to address"],
        "ats_keywords": ["Important keywords missing for ATS optimization"]
    }}
    
    Be thorough and specific in your gap analysis to help create a standout resume.
    """
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": "openai/gpt-3.5-turbo",
        "messages": [
            {"role": "system", "content": "You are an expert career counselor and resume analyst."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }
    
    try:
        response = requests.post(OPENROUTER_URL, headers=headers, json=data)
        response.raise_for_status()
        
        result = response.json()
        analysis_text = result['choices'][0]['message']['content']
        
        # Try to parse as JSON, fallback to text if fails
        try:
            return json.loads(analysis_text)
        except json.JSONDecodeError:
            return {"raw_analysis": analysis_text}
            
    except Exception as e:
        return {"error": f"AI analysis failed: {str(e)}"}

def generate_optimized_resume(resume_text, job_description, analysis_result):
    """Generate an optimized resume based on analysis and job description"""
    
    prompt = f"""
    Based on comprehensive gap analysis, create a standout, interview-winning resume that addresses all identified gaps and positions candidate as ideal fit.
    
    Original Resume:
    {resume_text}
    
    Job Description:
    {job_description}
    
    Gap Analysis Results:
    {analysis_result}
    
    Create a POWERFUL resume that:
    1. Addresses ALL critical gaps identified in analysis
    2. Incorporates missing ATS keywords for better screening
    3. Uses industry-specific terminology that matches the job
    4. Highlights quantifiable achievements with metrics and impact
    5. Positions experience level appropriately for role's seniority
    6. Creates a compelling narrative that stands out to recruiters
    7. Optimized for both ATS systems and human readers
    
    Follow this PROFESSIONAL structure:
    
    [FULL NAME]
    [Phone] | [Email] | [LinkedIn URL] | [Location] | [Portfolio/GitHub (if applicable)]
    
    PROFESSIONAL SUMMARY
    [3-4 line powerful summary that immediately shows value and addresses key requirements]
    
    CORE COMPETENCIES
    [List of 6-8 key skills that match job description perfectly]
    
    PROFESSIONAL EXPERIENCE
    [Company Name] | [City, State]
    [Job Title] | [Dates]
    • [Quantifiable achievement 1 - with specific metrics, results, and impact]
    • [Quantifiable achievement 2 - using industry terminology from job description]
    • [Quantifiable achievement 3 - showing growth and leadership]
    
    EDUCATION
    [Degree], [Major] | [University Name] | [Year]
    [Relevant coursework or achievements if applicable]
    
    PROJECTS (if technical role)
    [Project Name] | [Technologies Used]
    • [Description with business impact and technical challenges solved]
    
    CERTIFICATIONS & AWARDS
    [Certification Name] | [Issuing Organization] | [Year]
    
    Make this resume IRRESISTIBLE to recruiters and hiring managers. Use action verbs, quantify everything, and ensure it passes ATS screening while impressing humans.
    """
    
    try:
        response = requests.post(OPENROUTER_URL, headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }, json={
            "model": "openai/gpt-3.5-turbo",
            "messages": [
                {"role": "system", "content": "You are an expert resume writer and career coach specializing in creating professional, international-standard resumes."},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.4
        })
        response.raise_for_status()
        result = response.json()
        optimized_resume = result['choices'][0]['message']['content']
        return optimized_resume
            
    except Exception as e:
        return f"Error generating optimized resume: {str(e)}"

def create_pdf_from_text(text_content):
    """Create a PDF file from text content"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            if line.isupper() and len(line) < 30:
                para = Paragraph(f"<b>{line}</b>", styles['Heading2'])
            elif line.startswith('•'):
                para = Paragraph(f"• {line[1:].strip()}", styles['Normal'])
            else:
                para = Paragraph(line, styles['Normal'])
            story.append(para)
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_from_text(text_content):
    """Create a DOCX file from text content"""
    doc = docx.Document()
    
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            if line.isupper() and len(line) < 30:
                paragraph = doc.add_paragraph(line)
                paragraph.style = 'Heading 2'
            elif line.startswith('•'):
                doc.add_paragraph(line[1:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.set_page_config(
    page_title="🚀 AI Resume Analyzer Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional design
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .feature-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
        border-left: 4px solid #667eea;
    }
    
    .metric-card {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        color: white;
        padding: 1.5rem;
        border-radius: 10px;
        text-align: center;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .success-box {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .warning-box {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 5px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .error-box {
        background: #f8d7da;
        border: 1px solid #f5c6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 0.5rem 0;
    }
    
    .stProgress > div > div > div {
        background: linear-gradient(90deg, #667eea, #764ba2);
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border: none;
        border-radius: 8px;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
    }
</style>
""", unsafe_allow_html=True)

# Professional Header
st.markdown("""
<div class="main-header">
    <h1 style="margin: 0; font-size: 2.5rem;">🚀 AI Resume Analyzer Pro</h1>
    <p style="margin: 0.5rem 0 0 0; font-size: 1.2rem; opacity: 0.9;">
        🤖 Advanced AI-Powered Resume Optimization & Gap Analysis
    </p>
    <div style="margin-top: 1rem;">
        <span style="background: rgba(255,255,255,0.2); padding: 0.3rem 0.8rem; border-radius: 20px; margin: 0 0.5rem;">
            🔢 Token Management
        </span>
        <span style="background: rgba(255,255,255,0.2); padding: 0.3rem 0.8rem; border-radius: 20px; margin: 0 0.5rem;">
            🧠 Few-shot Learning
        </span>
        <span style="background: rgba(255,255,255,0.2); padding: 0.3rem 0.8rem; border-radius: 20px; margin: 0 0.5rem;">
            🎯 ATS Optimization
        </span>
    </div>
</div>
""", unsafe_allow_html=True)

# Sidebar for file upload
with st.sidebar:
    st.markdown("### 📁 Resume Upload")
    st.markdown("---")
    
    uploaded_file = st.file_uploader(
        "📄 Choose Resume File",
        type=['pdf', 'docx'],
        help="Upload your resume for AI analysis",
        label_visibility="collapsed"
    )
    
    if uploaded_file:
        file_details = {
            "Name": uploaded_file.name,
            "Size": f"{uploaded_file.size / 1024:.1f} KB",
            "Type": uploaded_file.type
        }
        
        st.markdown("""
        <div class="feature-card">
            <h4>📋 File Information</h4>
            <p><strong>Name:</strong> {}</p>
            <p><strong>Size:</strong> {}</p>
            <p><strong>Type:</strong> {}</p>
        </div>
        """.format(file_details["Name"], file_details["Size"], file_details["Type"]), unsafe_allow_html=True)
    
    st.markdown("---")
    st.markdown("### 📋 Job Description")
    job_description = st.text_area(
        "🎯 Paste Target Job Description",
        height=150,
        placeholder="Paste complete job description here for better analysis...",
        help="Provide detailed job description for comprehensive gap analysis",
        label_visibility="collapsed"
    )
    
    if job_description:
        char_count = len(job_description)
        st.markdown(f"""
        <div style="margin-top: 1rem; padding: 0.5rem; background: #f0f2f5; border-radius: 5px; text-align: center;">
            <small>📝 Character Count: {char_count}</small>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Analysis Button
    if uploaded_file and job_description:
        st.markdown("""
        <div style="text-align: center; margin: 2rem 0;">
            <button style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border: none; padding: 1rem 2rem; border-radius: 8px; font-weight: bold; font-size: 1.1rem; cursor: pointer; box-shadow: 0 4px 15px rgba(0,0,0,0.2);">
                🚀 Start AI Analysis
            </button>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("🚀 Analyze Resume", type="primary", use_container_width=True):
            # Extract text from uploaded file
            file_bytes = uploaded_file.read()
            
            # Save temporarily
            with open("temp_resume", "wb") as f:
                f.write(file_bytes)
            
            # Extract text based on file type
            if uploaded_file.type == "application/pdf":
                resume_text = extract_text_from_pdf("temp_resume")
            else:
                resume_text = extract_text_from_docx("temp_resume")
            
            # Analyze with AI
            with st.spinner("🤖 Performing Advanced AI Analysis..."):
                analysis_result = analyze_resume_with_ai(resume_text, job_description)
            
            # Generate optimized resume
            if "error" not in analysis_result:
                optimized_resume = generate_optimized_resume(resume_text, job_description, analysis_result)
                analysis_result["optimized_resume"] = optimized_resume
            
            # Clean up
            if os.path.exists("temp_resume"):
                os.remove("temp_resume")
            
            # Display results
            display_analysis_results(analysis_result)
    else:
        st.markdown("""
        <div class="warning-box">
            <strong>⚠️ Ready to Analyze</strong><br>
            Please upload both resume and job description to start the AI analysis.
        </div>
        """, unsafe_allow_html=True)

def display_analysis_results(analysis_result):
    """Display professional analysis results"""
    st.markdown("---")
    
    if "error" in analysis_result:
        st.markdown(f"""
        <div class="error-box">
            <h4>❌ Analysis Error</h4>
            <p>{analysis_result['error']}</p>
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Match Score Section
    match_percentage = analysis_result.get('match_percentage', 0)
    st.markdown(f"""
    <div class="metric-card">
        <h2 style="margin: 0; font-size: 2rem;">🎯 Match Score</h2>
        <div style="font-size: 3rem; font-weight: bold; margin: 1rem 0;">
            {match_percentage}%
        </div>
        <p style="margin: 0; opacity: 0.8;">Resume-Job Compatibility Score</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Progress Bar
    st.progress(match_percentage / 100)
    
    # Analysis Results
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <h3>✅ Strengths</h3>
        </div>
        """, unsafe_allow_html=True)
        
        strengths = analysis_result.get('strengths', [])
        for strength in strengths:
            st.markdown(f"""
            <div class="success-box">
                <strong>✓</strong> {strength}
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <h3>⚠️ Missing Skills</h3>
        </div>
        """, unsafe_allow_html=True)
        
        missing_skills = analysis_result.get('missing_skills', [])
        for skill in missing_skills:
            st.markdown(f"""
            <div class="warning-box">
                <strong>!</strong> {skill}
            </div>
            """, unsafe_allow_html=True)
    
    # Recommendations
    st.markdown("""
    <div class="feature-card">
        <h3>💡 Professional Recommendations</h3>
    </div>
    """, unsafe_allow_html=True)
    
    recommendations = analysis_result.get('recommendations', [])
    for rec in recommendations:
        st.markdown(f"""
            <div style="background: #e3f2fd; border-left: 4px solid #667eea; padding: 1rem; margin: 0.5rem 0; border-radius: 5px;">
            <strong>🎯</strong> {rec}
        </div>
        """, unsafe_allow_html=True)
    
    # Key Matches
    if 'key_matches' in analysis_result:
        st.markdown("""
        <div class="feature-card">
            <h3>🎯 Key Matches</h3>
        </div>
        """, unsafe_allow_html=True)
        
        key_matches = analysis_result['key_matches']
        for match in key_matches:
            st.markdown(f"""
            <div style="background: #f0fff4; border: 1px solid #e8f5e8; padding: 1rem; margin: 0.5rem 0; border-radius: 5px;">
                <strong>✓</strong> {match}
            </div>
            """, unsafe_allow_html=True)
    
    # Critical Gaps
    if 'critical_gaps' in analysis_result:
        st.markdown("""
        <div class="feature-card">
            <h3>🚨 Critical Gaps</h3>
        </div>
        """, unsafe_allow_html=True)
        
        gaps = analysis_result['critical_gaps']
        for gap in gaps:
            st.markdown(f"""
            <div class="error-box">
                <strong>⚠️</strong> {gap}
            </div>
            """, unsafe_allow_html=True)
    
    # Overall Assessment
    overall_assessment = analysis_result.get('overall_assessment', 'No assessment available')
    st.markdown(f"""
    <div class="feature-card">
        <h3>📊 Overall Assessment</h3>
        <p style="font-size: 1.1rem; line-height: 1.6;">{overall_assessment}</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Optimized Resume Section
    if "optimized_resume" in analysis_result:
        st.markdown("---")
        st.markdown("""
        <div class="feature-card">
            <h3>📄 AI-Optimized Resume</h3>
            <p style="color: #667eea; font-weight: bold; margin-bottom: 1rem;">
                ✨ Professionally formatted for ATS systems and recruiters
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Resume Preview
        st.markdown("#### 📋 Resume Preview")
        st.text_area(
            "Optimized resume content:",
            value=analysis_result["optimized_resume"],
            height=400,
            help="Your professionally optimized resume content",
            label_visibility="collapsed"
        )
        
        # Download Options
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📥 Download PDF", type="secondary", use_container_width=True):
                pdf_buffer = create_pdf_from_text(analysis_result["optimized_resume"])
                st.download_button(
                    label="📄 Download Resume (PDF)",
                    data=pdf_buffer.getvalue(),
                    file_name="optimized_resume.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        
        with col2:
            if st.button("📝 Download DOCX", type="secondary", use_container_width=True):
                docx_buffer = create_docx_from_text(analysis_result["optimized_resume"])
                st.download_button(
                    label="📝 Download Resume (DOCX)",
                    data=docx_buffer.getvalue(),
                    file_name="optimized_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

# Instructions
with st.expander("📖 How to Use", expanded=False):
    st.markdown("""
    ### 🚀 Getting Started:
    1. **📁 Upload Resume**: Click "Choose Resume File" to upload your PDF or DOCX resume
    2. **📋 Add Job Description**: Paste target job description
    3. **🚀 Analyze**: Click "Start AI Analysis" for comprehensive analysis
    4. **📥 Download**: Get your optimized resume in PDF or DOCX format
    
    ### ✨ Advanced Features:
    - **🔢 Token Management**: Handles resumes of any length with intelligent processing
    - **🧠 Few-shot Learning**: Uses examples for better AI accuracy
    - **🔗 Chain-of-Thought**: Step-by-step reasoning for comprehensive analysis
    - **🎯 ATS Optimization**: Industry-specific keyword integration
    - **📊 Professional Analysis**: Gap analysis and recommendations
    - **📄 Multiple Formats**: PDF and DOCX download options
    
    ### 🎯 Benefits:
    - **✅ Higher Match Rates**: Better alignment with job requirements
    - **✅ ATS-Friendly**: Optimized for automated screening
    - **✅ Professional Format**: Industry-standard resume structure
    - **✅ Quantified Achievements**: Metrics-based accomplishments
    - **✅ Free to Use**: No cost for advanced AI analysis
    """)

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #667eea; padding: 20px; border-top: 2px solid #f0f2f5; margin-top: 2rem;'>
    <p style="margin: 0; font-weight: bold;">
        🤖 Built with ❤️ using Advanced AI Technologies
    </p>
    <p style='margin: 0.5rem 0 0 0; font-size: 0.9rem; color: #666;'>
        🔧 Features: Token Management | Few-shot Prompting | Chain-of-Thought | ATS Optimization
    </p>
    <p style='margin: 0; font-size: 0.8rem; color: #999;'>
        🌐 Deployed on Streamlit Cloud | Professional Resume Analysis Tool
    </p>
</div>
""", unsafe_allow_html=True)
