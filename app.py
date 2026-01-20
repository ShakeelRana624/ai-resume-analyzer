from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file
import os
import PyPDF2
import docx
import json
from werkzeug.utils import secure_filename
import requests
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from io import BytesIO
import tempfile

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# OpenRouter API configuration
OPENROUTER_API_KEY = "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"  # OpenRouter API key
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_resume_text(file_path, file_extension):
    """Extract text from resume file based on its type"""
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")

def generate_optimized_resume(resume_text, job_description, analysis_result):
    """Generate an optimized resume based on analysis and job description"""
    
    prompt = f"""
    Based on the comprehensive gap analysis, create a standout, interview-winning resume that addresses all identified gaps and positions the candidate as the ideal fit.
    
    Original Resume:
    {resume_text}
    
    Job Description:
    {job_description}
    
    Gap Analysis Results:
    {analysis_result}
    
    Create a POWERFUL resume that:
    1. Addresses ALL critical gaps identified in the analysis
    2. Incorporates missing ATS keywords for better screening
    3. Uses industry-specific terminology that matches the job
    4. Highlights quantifiable achievements with metrics and impact
    5. Positions experience level appropriately for the role's seniority
    6. Creates a compelling narrative that stands out to recruiters
    7. Optimized for both ATS systems and human readers
    
    Follow this PROFESSIONAL structure:
    
    [FULL NAME]
    [Phone] | [Email] | [LinkedIn URL] | [Location] | [Portfolio/GitHub (if applicable)]
    
    PROFESSIONAL SUMMARY
    [3-4 line powerful summary that immediately shows value and addresses key requirements]
    
    CORE COMPETENCIES
    [List of 6-8 key skills that match job description perfectly]
    
    PROFESSIONAL EXPERIENCE
    [Company Name] | [City, State]
    [Job Title] | [Dates]
    • [Quantifiable achievement 1 - with specific metrics, results, and impact]
    • [Quantifiable achievement 2 - using industry terminology from job description]
    • [Quantifiable achievement 3 - showing growth and leadership]
    
    EDUCATION
    [Degree], [Major] | [University Name] | [Year]
    [Relevant coursework or achievements if applicable]
    
    PROJECTS (if technical role)
    [Project Name] | [Technologies Used]
    • [Description with business impact and technical challenges solved]
    
    CERTIFICATIONS & AWARDS
    [Certification Name] | [Issuing Organization] | [Year]
    
    Make this resume IRRESISTIBLE to recruiters and hiring managers. Use action verbs, quantify everything, and ensure it passes ATS screening while impressing humans.
    """
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": "openai/gpt-3.5-turbo",
        "messages": [
            {"role": "system", "content": "You are an expert resume writer and career coach specializing in creating professional, international-standard resumes."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.4
    }
    
    try:
        response = requests.post(OPENROUTER_URL, headers=headers, json=data)
        response.raise_for_status()
        
        result = response.json()
        optimized_resume = result['choices'][0]['message']['content']
        return optimized_resume
            
    except Exception as e:
        return f"Error generating optimized resume: {str(e)}"

import tiktoken

def count_tokens(text):
    """Count tokens in text using tiktoken"""
    try:
        encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        return len(encoding.encode(text))
    except:
        # Fallback: rough estimation (1 token ≈ 4 characters)
        return len(text) // 4

def chunk_text(text, max_tokens=3000):
    """Split text into chunks that don't exceed max tokens"""
    if count_tokens(text) <= max_tokens:
        return [text]
    
    chunks = []
    current_chunk = ""
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    
    for paragraph in paragraphs:
        test_chunk = current_chunk + paragraph + '\n\n' if current_chunk else paragraph
        
        if count_tokens(test_chunk) <= max_tokens:
            current_chunk = test_chunk
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            # Split long paragraph
            sentences = paragraph.split('. ')
            for sentence in sentences:
                test_sentence = current_chunk + sentence + '. ' if current_chunk else sentence + '. '
                if count_tokens(test_sentence) <= max_tokens:
                    current_chunk = test_sentence
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + '. '
            
            current_chunk = ""
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

def analyze_resume_with_ai(resume_text, job_description):
    """Analyze resume against job description using OpenRouter API"""
    
    # Check token count and chunk if necessary
    total_tokens = count_tokens(resume_text)
    print(f"Resume token count: {total_tokens}")
    
    if total_tokens > 3000:
        print("Resume too long, chunking for analysis...")
        resume_chunks = chunk_text(resume_text, max_tokens=2500)
        
        # Analyze first chunk for basic info
        first_chunk_prompt = f"""
        Please analyze this resume chunk against the job description and provide initial assessment.
        
        Job Description:
        {job_description}
        
        Resume Chunk (Part 1 of {len(resume_chunks)}):
        {resume_chunks[0]}
        
        This is part of a longer resume. Focus on:
        1. Basic candidate information
        2. Core skills and experience level
        3. Education background
        4. Initial gap analysis
        
        Provide JSON response with:
        {{
            "match_percentage": 75,
            "strengths": ["Initial strengths from this chunk"],
            "missing_skills": ["Skills not found in this chunk"],
            "recommendations": ["Initial recommendations"],
            "key_matches": ["Matches found in this chunk"],
            "experience_alignment": "Experience level assessment",
            "overall_assessment": "Initial assessment",
            "critical_gaps": ["Gaps identified in this chunk"],
            "ats_keywords": ["Keywords from job description"],
            "needs_full_analysis": true
        }}
        """
        
        try:
            response = requests.post(OPENROUTER_URL, headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json"
            }, json={
                "model": "openai/gpt-3.5-turbo",
                "messages": [
                    {"role": "system", "content": "You are an expert resume analyst conducting initial analysis of a long resume."},
                    {"role": "user", "content": first_chunk_prompt}
                ],
                "temperature": 0.3
            })
            response.raise_for_status()
            result = response.json()
            return json.loads(result['choices'][0]['message']['content'])
            
        except Exception as e:
            return {"error": f"Initial analysis failed: {str(e)}"}
    
    # For shorter resumes, use regular analysis with Few-shot prompting
    prompt = f"""
    You are an expert resume analyst with 10+ years of experience in career counseling and ATS optimization.
    
    Your task is to analyze a resume against a job description using Chain-of-Thought reasoning.
    
    Here are examples of high-quality analysis:
    
    Example 1 - Strong Match:
    {{
        "match_percentage": 85,
        "strengths": ["5+ years Python development", "AWS certified", "Led team of 5"],
        "missing_skills": ["No critical gaps identified"],
        "recommendations": ["Add quantifiable metrics to achievements"],
        "critical_gaps": ["Minor: Add project outcomes"],
        "ats_keywords": ["All key terms present"]
    }}
    
    Example 2 - Medium Match:
    {{
        "match_percentage": 65,
        "strengths": ["Strong technical background", "Good communication skills"],
        "missing_skills": ["Cloud experience", "Leadership examples"],
        "recommendations": ["Add cloud certifications", "Show leadership examples"],
        "critical_gaps": ["Missing cloud architecture experience"],
        "ats_keywords": ["Cloud, Kubernetes missing"]
    }}
    
    Now analyze this resume using the same thorough approach:
    
    Job Description:
    {job_description}
    
    Resume:
    {resume_text}
    
    Chain-of-Thought Process:
    1. First, identify all technical skills in the job description
    2. Then, map each skill to the resume content
    3. Next, assess experience level alignment
    4. Finally, identify critical gaps and ATS optimization opportunities
    
    Focus on identifying:
    1. Critical gaps between resume requirements and job needs
    2. Missing keywords and skills that ATS systems look for
    3. Experience level alignment with seniority requirements
    4. Industry-specific terminology gaps
    5. Quantifiable achievements that are missing
    
    Please provide analysis in the following JSON format:
    {{
        "match_percentage": 85,
        "strengths": ["List of key strengths found in resume"],
        "missing_skills": ["List of important skills missing from resume"],
        "recommendations": ["List of specific recommendations to improve resume"],
        "key_matches": ["List of good matches between resume and job description"],
        "experience_alignment": "Assessment of how well experience aligns with requirements",
        "overall_assessment": "Brief overall assessment",
        "critical_gaps": ["List of most critical gaps to address"],
        "ats_keywords": ["Important keywords missing for ATS optimization"]
    }}
    
    Be thorough and specific in your gap analysis to help create a standout resume.
    """
    
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    
    data = {
        "model": "openai/gpt-3.5-turbo",
        "messages": [
            {"role": "system", "content": "You are an expert career counselor and resume analyst."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3
    }
    
    try:
        response = requests.post(OPENROUTER_URL, headers=headers, json=data)
        response.raise_for_status()
        
        result = response.json()
        analysis_text = result['choices'][0]['message']['content']
        
        # Try to parse as JSON, fallback to text if fails
        try:
            return json.loads(analysis_text)
        except json.JSONDecodeError:
            return {"raw_analysis": analysis_text}
            
    except Exception as e:
        return {"error": f"AI analysis failed: {str(e)}"}

def create_pdf_from_text(text_content):
    """Create a PDF file from text content"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Split text into lines and create paragraphs
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            # Handle different formatting for headers
            if line.isupper() and len(line) < 30:
                para = Paragraph(f"<b>{line}</b>", styles['Heading2'])
            elif line.startswith('•'):
                para = Paragraph(f"• {line[1:].strip()}", styles['Normal'])
            else:
                para = Paragraph(line, styles['Normal'])
            story.append(para)
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_docx_from_text(text_content):
    """Create a DOCX file from text content"""
    doc = docx.Document()
    
    # Split text into lines and add to document
    lines = text_content.split('\n')
    for line in lines:
        if line.strip():
            # Handle different formatting for headers
            if line.isupper() and len(line) < 30:
                paragraph = doc.add_paragraph(line)
                paragraph.style = 'Heading 2'
            elif line.startswith('•'):
                doc.add_paragraph(line[1:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/download/pdf')
def download_pdf():
    """Download optimized resume as PDF"""
    resume_text = request.args.get('resume', '')
    if not resume_text:
        return jsonify({"error": "No resume content provided"}), 400
    
    try:
        pdf_buffer = create_pdf_from_text(resume_text)
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name='optimized_resume.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500

@app.route('/download/docx')
def download_docx():
    """Download optimized resume as DOCX"""
    resume_text = request.args.get('resume', '')
    if not resume_text:
        return jsonify({"error": "No resume content provided"}), 400
    
    try:
        docx_buffer = create_docx_from_text(resume_text)
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name='optimized_resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({"error": f"DOCX generation failed: {str(e)}"}), 500

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        # Check if file was uploaded
        if 'resume' not in request.files:
            return jsonify({"error": "No resume file uploaded"}), 400
        
        file = request.files['resume']
        job_description = request.form.get('job_description', '')
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
            
        if not allowed_file(file.filename):
            return jsonify({"error": "Invalid file type. Please upload PDF or DOCX files only"}), 400
            
        if not job_description.strip():
            return jsonify({"error": "Job description is required"}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Extract text from resume
        file_extension = filename.rsplit('.', 1)[1].lower()
        resume_text = extract_resume_text(file_path, file_extension)
        
        if not resume_text.strip():
            return jsonify({"error": "Could not extract text from the uploaded file"}), 400
        
        # Analyze with AI (no API key needed from user)
        analysis_result = analyze_resume_with_ai(resume_text, job_description)
        
        # Generate optimized resume if analysis was successful
        if "error" not in analysis_result:
            optimized_resume = generate_optimized_resume(resume_text, job_description, analysis_result)
            analysis_result["optimized_resume"] = optimized_resume
        
        # Clean up uploaded file
        os.remove(file_path)
        
        return jsonify(analysis_result)
        
    except Exception as e:
        return jsonify({"error": f"Analysis failed: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
